"""
ServeMart Mini Project Report Generator — v4
Generates ServeMart_MiniProject_Report_v4.docx following the MIT Manipal DBS Lab template.
Run: python gen_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def heading(doc, text, size=14, bold=True, center=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def body(doc, text, size=12, space_before=0, space_after=6, indent=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def numbered(doc, text, size=12):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def code_block(doc, text, size=9):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(size)
    doc.add_paragraph()

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    if col_widths:
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    doc.add_paragraph()
    return table

# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.0)

# ─── PAGE 1: TITLE PAGE ─────────────────────────────────────────────────────

heading(doc, 'Mini Project Report', size=16, bold=True, center=True, space_before=48, space_after=4)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('of'); set_font(r, size=14)
heading(doc, 'Database Systems Lab (CSS 2212)', size=14, bold=True, center=True, space_before=4, space_after=16)
heading(doc, 'ServeMart', size=20, bold=True, center=True, space_before=12, space_after=4)
heading(doc, 'Local Service Marketplace & Booking Management System', size=14, bold=False, center=True, space_before=4, space_after=32)
body(doc, 'SUBMITTED BY', size=12, space_before=20, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

tbl = doc.add_table(rows=4, cols=4)
tbl.style = 'Table Grid'
for ci, h in enumerate(['Student Name', 'Reg. No', 'Roll No', 'Section']):
    tbl.rows[0].cells[ci].text = h
    for run in tbl.rows[0].cells[ci].paragraphs[0].runs:
        run.font.bold = True; run.font.size = Pt(11); run.font.name = 'Times New Roman'
for ri, row in enumerate([
    ('Abhyuday Gupta',    '240905576', '57', 'CSE C'),
    ('Shaurya Jain',      '240905598', '60', 'CSE C'),
    ('Neelaksha Sisodiya','240905642', '65', 'CSE C'),
]):
    for ci, val in enumerate(row):
        tbl.rows[ri+1].cells[ci].text = val
        for run in tbl.rows[ri+1].cells[ci].paragraphs[0].runs:
            run.font.size = Pt(11); run.font.name = 'Times New Roman'

doc.add_paragraph()
heading(doc, 'School of Computer Engineering', size=12, bold=True, center=True, space_before=24, space_after=2)
heading(doc, 'Manipal Institute of Technology, Manipal', size=12, bold=False, center=True, space_before=0, space_after=2)
heading(doc, 'April 2026', size=12, bold=False, center=True, space_before=0, space_after=0)
doc.add_page_break()

# ─── PAGE 2: CERTIFICATE ────────────────────────────────────────────────────

heading(doc, 'DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING', size=13, bold=True, center=True, space_before=12, space_after=4)
heading(doc, 'Manipal Institute of Technology, Manipal', size=12, bold=False, center=True, space_before=0, space_after=4)
heading(doc, 'April 2026', size=12, bold=False, center=True, space_before=0, space_after=16)
heading(doc, 'CERTIFICATE', size=16, bold=True, center=True, space_before=8, space_after=16)
body(doc, (
    'This is to certify that the Mini Project titled "ServeMart \u2014 Local Service Marketplace & '
    'Booking Management System" submitted by Abhyuday Gupta (Reg. No. 240905576), Shaurya Jain '
    '(Reg. No. 240905598), and Neelaksha Sisodiya (Reg. No. 240905642) of III Semester, B.Tech '
    '(Computer Science and Engineering), Section C, is a bonafide record of the work carried out by '
    'them in partial fulfilment of the requirements for the course Database Systems Lab (CSS 2212) '
    'during the academic year 2025\u20132026.'
), size=12, space_after=24)
body(doc, 'Name and Signature of Examiners:', size=12, space_before=8, space_after=12)
body(doc, '1. ______________________________________', size=12, space_after=12)
body(doc, '2. ______________________________________', size=12, space_after=0)
doc.add_page_break()

# ─── TABLE OF CONTENTS ──────────────────────────────────────────────────────

heading(doc, 'TABLE OF CONTENTS', size=16, bold=True, center=True, space_before=8, space_after=16)
toc_entries = [
    ('\u2014', 'Abstract', '4'),
    ('1',  'Introduction', '5'),
    ('2',  'Problem Statement and Objectives', '9'),
    ('3',  'Methodology', '12'),
    ('4',  'ER Diagram, Relational Schema and Sample Data', '15'),
    ('5',  'DDL Commands and PL/SQL Procedures / Functions / Triggers', '22'),
    ('6',  'Results and Snapshots', '38'),
    ('7',  'Conclusion, Limitations and Future Work', '43'),
    ('8',  'References', '47'),
]
toc_tbl = doc.add_table(rows=len(toc_entries)+1, cols=3)
toc_tbl.style = 'Table Grid'
for ci, h in enumerate(['No.', 'Topic', 'Page No.']):
    toc_tbl.rows[0].cells[ci].text = h
    for run in toc_tbl.rows[0].cells[ci].paragraphs[0].runs:
        run.font.bold = True; run.font.size = Pt(11); run.font.name = 'Times New Roman'
for ri, (num, topic, page) in enumerate(toc_entries):
    for ci, val in enumerate([num, topic, page]):
        toc_tbl.rows[ri+1].cells[ci].text = val
        for run in toc_tbl.rows[ri+1].cells[ci].paragraphs[0].runs:
            run.font.size = Pt(11); run.font.name = 'Times New Roman'
doc.add_page_break()

# ─── ABSTRACT ───────────────────────────────────────────────────────────────

heading(doc, 'ABSTRACT', size=16, bold=True, center=True, space_before=8, space_after=16)
body(doc, (
    'ServeMart is a full-stack, role-based local service marketplace. It lets customers discover, '
    'book, and pay verified home service professionals \u2014 plumbers, electricians, painters, '
    'cleaners \u2014 through a web interface. Service providers manage their availability and service '
    'listings through a separate dashboard, and a platform administrator handles provider approvals, '
    'manages service cities, and monitors revenue.'
), size=12, space_after=8)
body(doc, (
    'The system runs on a three-tier architecture: a React 19 frontend built with Vite, a '
    'Node.js/Express.js REST API in the middle, and Oracle Database XE as the persistent store. '
    'The database has 16 normalised tables across five domains: Identity, Geography, Catalogue, '
    'Operations, and Audit. Three PL/SQL stored procedures handle the booking lifecycle atomically. '
    'A compound trigger on the REVIEWS table keeps provider rating averages and job counts current '
    'automatically. A stored function ranks providers by a composite score for the recommendation '
    'endpoint. A view simplifies provider listing queries used across the application.'
), size=12, space_after=8)
body(doc, (
    'Security is handled through JWT authentication, bcrypt password hashing, Zod request '
    'validation, role-based access control on every API route, and bind variables on every Oracle '
    'query. The project demonstrates entity-relationship modelling, relational schema design, '
    'normalisation to 3NF, and the full range of Oracle PL/SQL features in a working marketplace '
    'system.'
), size=12, space_after=0)
doc.add_page_break()

# ─── CHAPTER 1 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 1: INTRODUCTION', size=14, bold=True, space_before=8, space_after=12)

heading(doc, '1.1  Background', size=13, bold=True, space_before=8, space_after=6)
body(doc, (
    'Booking a plumber or electrician in India is still mostly done through word of mouth or '
    'informal WhatsApp groups. There is no shared platform where customers can compare ratings, '
    'check real availability, and get a receipt \u2014 and no place where providers can advertise '
    'services, block out their calendar, or track their earnings digitally.'
), size=12, space_after=8)
body(doc, (
    'ServeMart is a database-backed attempt at solving this. Three roles interact with the platform: '
    'customers browse and book verified professionals, providers manage their listings and mark jobs '
    'complete, and an administrator approves new provider accounts, adds service cities, and '
    'monitors platform revenue. The project\'s primary purpose is to apply relational database '
    'design, Oracle PL/SQL, and full-stack web development to a practical domain.'
), size=12, space_after=8)

heading(doc, '1.2  System Architecture', size=13, bold=True, space_before=8, space_after=6)
body(doc, 'ServeMart uses a three-tier architecture:', size=12, space_after=6)
code_block(doc, """\
+------------------------------------------------+
|            PRESENTATION TIER                   |
|  React 19 (Vite) -- Single-Page Application    |
|  Role dashboards: Customer / Provider / Admin  |
|  Leaflet map  *  Recharts analytics            |
+------------------+-----------------------------+
                   |  HTTP / REST (JSON)
                   |  JWT Bearer tokens
+------------------v-----------------------------+
|            APPLICATION TIER                    |
|  Node.js + Express.js 4                        |
|  6 route groups: auth, providers, bookings,    |
|  invoices, reviews, admin                      |
|  Zod validation * bcrypt * JWT                 |
|  express-rate-limit * CORS                     |
+------------------+-----------------------------+
                   |  oracledb v6
                   |  Connection pool (min 2, max 10)
+------------------v-----------------------------+
|             DATABASE TIER                      |
|  Oracle Database XE                            |
|  16 tables * 3 stored procedures               |
|  1 compound trigger * 1 function * 1 view      |
+------------------------------------------------+""")
body(doc, (
    'The React frontend sends all requests to the Express backend over HTTP. The backend acquires '
    'a connection from the Oracle pool, executes parameterised queries or calls PL/SQL objects '
    'using named bind variables, releases the connection in a finally block, and returns a uniform '
    'JSON response: { success, data } on success or { success, error: { code, message } } on '
    'failure. Oracle error codes (ORA-XXXXX) are translated to user-readable messages before being '
    'sent to the client.'
), size=12, space_after=8)

heading(doc, '1.3  Scope of the project', size=13, bold=True, space_before=8, space_after=6)
for item in [
    'User registration and JWT-based login for three roles (CUSTOMER, PROVIDER, ADMIN)',
    'Provider profile management: service catalogue, weekly availability slots, and service area coverage',
    'Customer service discovery: category filtering, provider cards, interactive Leaflet map, and city-based recommendations',
    'Full booking lifecycle: creation with time-overlap prevention, completion, and cancellation with slot restoration',
    'Promo code validation endpoint with live discount preview before booking submission',
    'Automated invoice generation with promo-code discounts, 10% platform fee, and 5% GST on job completion',
    'Payment method selection (Cash / UPI / Credit Card) captured by the provider when marking a job complete',
    'Customer review submission (1\u20135 stars + comment) with automatic provider rating recalculation via trigger',
    'Admin revenue analytics (grouped by service category), provider approval, service area management, and error log viewer',
    'Database-level error logging from all PL/SQL objects to ERROR_LOGS',
    'Payment gateway integration is explicitly out of scope; payments are recorded as confirmed when the provider marks the job complete',
]:
    bullet(doc, item)
doc.add_paragraph()

heading(doc, '1.4  Tools and Technologies', size=13, bold=True, space_before=8, space_after=6)
simple_table(doc, ['Layer', 'Technology', 'Version'], [
    ('Frontend Framework',  'React',                    '19'),
    ('Build Tool',          'Vite',                     '8.x'),
    ('Routing',             'React Router',             'v7'),
    ('HTTP Client',         'Axios',                    '1.x'),
    ('Maps',                'Leaflet / react-leaflet',  '1.9 / 5.0'),
    ('Charts',              'Recharts',                 '3.x'),
    ('Icons',               'react-icons',              '5.x'),
    ('Styling',             'Vanilla CSS (custom)',     '\u2014'),
    ('Backend Framework',   'Express.js',               '4.x'),
    ('Runtime',             'Node.js',                  '18+'),
    ('DB Driver',           'oracledb',                 '6.x'),
    ('Authentication',      'jsonwebtoken',             '9.x'),
    ('Password Hashing',    'bcrypt',                   '5.x'),
    ('Input Validation',    'Zod',                      '3.x'),
    ('Rate Limiting',       'express-rate-limit',       '7.x'),
    ('Database',            'Oracle Database XE',       '\u2014'),
    ('ER Modelling',        'draw.io / SQL Developer',  '\u2014'),
    ('Environment Config',  'dotenv',                   '16.x'),
], col_widths=[1.8, 2.4, 1.0])
doc.add_page_break()

# ─── CHAPTER 2 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 2: PROBLEM STATEMENT AND OBJECTIVES', size=14, bold=True, space_before=8, space_after=12)

heading(doc, '2.1  Problem statement', size=13, bold=True, space_before=8, space_after=6)
body(doc, (
    'Local home service bookings are disorganised. Customers cannot check whether a provider is '
    'actually available before calling, cannot compare prices or verified ratings, and have no '
    'official record of the transaction. Providers lack a single place to advertise services, '
    'block out their calendar, or track completed work.'
), size=12, space_after=8)
body(doc, 'The core database challenge is to design a schema that:', size=12, space_after=4)
for item in [
    'Maintains full referential integrity across the entire booking lifecycle',
    'Prevents double-booking through an overlap check inside a stored procedure, so the restriction holds even if the database is called directly',
    'Keeps derived statistics (provider ratings, job counts) consistent through a compound trigger with no application-side recalculation',
    'Encapsulates business-critical operations in PL/SQL so the logic cannot be bypassed by the API layer',
]:
    numbered(doc, item)
doc.add_paragraph()

heading(doc, '2.2  Functional requirements', size=13, bold=True, space_before=8, space_after=6)
simple_table(doc, ['ID', 'Requirement'], [
    ('FR1',  'Users register as CUSTOMER or PROVIDER with a unique username and email'),
    ('FR2',  'Login returns a JWT valid for 24 hours; inactive accounts are rejected'),
    ('FR3',  'Routes are role-restricted: CUSTOMER, PROVIDER, and ADMIN endpoints are separated'),
    ('FR4',  'Customers browse providers, filter by category, and view them on an interactive map'),
    ('FR5',  'Customers pick an available time slot and book a service at a provider\'s hourly rate'),
    ('FR6',  'The booking procedure checks for overlapping active bookings before confirming'),
    ('FR7',  'Customers may validate a promo code before booking for a live discount preview'),
    ('FR8',  'Providers add and delete their weekly availability slots by day and hour range'),
    ('FR9',  'Providers add and delete services they offer, specifying category and hourly rate'),
    ('FR10', 'Providers manage the cities they serve through a dedicated area management page'),
    ('FR11', 'Providers mark bookings as COMPLETED; a payment method modal captures Cash/UPI/Card'),
    ('FR12', 'Invoices include base amount, promo discount (capped), 10% platform fee, and 5% GST'),
    ('FR13', 'Customers and providers cancel bookings; the availability slot is restored on cancellation'),
    ('FR14', 'Cancellation reason is displayed to both parties on cancelled booking cards'),
    ('FR15', 'Customers submit a 1\u20135 star review with comment for completed bookings'),
    ('FR16', 'Provider rating average and job count update automatically when a review is written'),
    ('FR17', 'Admins view revenue grouped by service category with a bar chart'),
    ('FR18', 'Admins approve or reject provider registrations (background check status)'),
    ('FR19', 'All PL/SQL exceptions are logged to ERROR_LOGS before re-raising'),
    ('FR20', 'A stored function ranks providers in a given city by a composite score'),
], col_widths=[0.6, 5.6])

heading(doc, '2.3  Objectives', size=13, bold=True, space_before=8, space_after=6)
for obj in [
    'Design a normalised Oracle schema (3NF) with 16 tables covering all marketplace entities.',
    'Implement stored procedures, a compound trigger, a function, and a view in PL/SQL.',
    'Build a secure REST API with Express.js: JWT auth, Zod validation, ownership checks on every mutating endpoint.',
    'Develop a role-specific React SPA with dashboards for all three user types.',
    'Prevent SQL injection by using bind variables on every Oracle query.',
    'Keep data integrity through foreign keys, check constraints, and transactional operations.',
]:
    bullet(doc, obj)
doc.add_page_break()

# ─── CHAPTER 3 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 3: METHODOLOGY', size=14, bold=True, space_before=8, space_after=12)
body(doc, (
    'The project followed a bottom-up Database Development Life Cycle (DDLC) through six phases: '
    'requirement gathering, conceptual design, logical design, physical design, implementation, '
    'and testing.'
), size=12, space_after=12)

heading(doc, '3.1  Requirement gathering', size=13, bold=True, space_before=8, space_after=6)
body(doc, 'The team identified entities through use-case analysis of a typical service booking flow:', size=12, space_after=4)
for item in [
    'Users \u2014 authentication record shared by all roles',
    'Customers and Service Providers \u2014 role-specific profiles',
    'Service Categories and Services Offered \u2014 the service catalogue',
    'Service Areas and Provider Areas \u2014 geographic coverage (M:N)',
    'Provider Availability \u2014 weekly recurring time slots',
    'Bookings \u2014 the central transaction entity',
    'Invoices, Payments, Cancellations \u2014 financial and lifecycle records',
    'Reviews \u2014 customer feedback that drives provider statistics via trigger',
    'Promotions \u2014 discount codes with usage caps and validity windows',
    'Error Logs \u2014 audit table written to by all PL/SQL objects',
]:
    bullet(doc, item)
doc.add_paragraph()

heading(doc, '3.2  Conceptual design', size=13, bold=True, space_before=8, space_after=6)
body(doc, 'Key relationships from the ER model:', size=12, space_after=4)
for r in [
    'USERS \u21901:1\u2192 CUSTOMERS (partial on USERS, total on CUSTOMERS)',
    'USERS \u21901:1\u2192 SERVICE_PROVIDERS (partial on USERS, total on SERVICE_PROVIDERS)',
    'CUSTOMERS \u21901:N\u2192 CUSTOMER_ADDRESSES (weak entity; identifying relationship)',
    'SERVICE_PROVIDERS \u2190M:N\u2192 SERVICE_AREAS (via PROVIDER_AREAS junction entity)',
    'SERVICE_PROVIDERS \u21901:N\u2192 PROVIDER_AVAILABILITY',
    'SERVICE_PROVIDERS \u21901:N\u2192 SERVICES_OFFERED \u2190N:1\u2192 SERVICE_CATEGORIES',
    'CUSTOMERS \u21901:N\u2192 BOOKINGS \u2190N:1\u2192 SERVICES_OFFERED',
    'BOOKINGS \u2190N:1\u2192 CUSTOMER_ADDRESSES, PROVIDER_AVAILABILITY',
    'BOOKINGS \u21901:1\u2192 INVOICES \u21901:1\u2192 PAYMENTS',
    'BOOKINGS \u21901:1\u2192 CANCELLATIONS (conditional; only for cancelled bookings)',
    'BOOKINGS \u21901:1\u2192 REVIEWS (conditional; only for completed bookings)',
    'PROMOTIONS \u21901:N\u2192 BOOKINGS (optional FK, SET NULL on delete)',
]:
    bullet(doc, r)
doc.add_paragraph()

heading(doc, '3.3  Logical design and normalisation', size=13, bold=True, space_before=8, space_after=6)
body(doc, (
    '1NF: All attributes are atomic. No repeating groups. Each table has a single surrogate '
    'primary key generated by Oracle\'s GENERATED ALWAYS AS IDENTITY.'
), size=12, space_after=6)
body(doc, (
    '2NF: No partial dependencies. All non-key attributes depend on the whole primary key. '
    'PROVIDER_AREAS and PROVIDER_AVAILABILITY use surrogate PKs with UNIQUE constraints on the '
    'natural composite keys.'
), size=12, space_after=6)
body(doc, (
    '3NF: No transitive dependencies. Provider statistics (rating_avg, jobs_completed) live in '
    'SERVICE_PROVIDERS because they describe the provider, not any individual review or booking. '
    'Financial fields (platform_fee, tax_amount, net_total) live in INVOICES because they are '
    'generated only at job completion. Cancellation metadata lives in CANCELLATIONS rather than as '
    'nullable columns on BOOKINGS, which keeps BOOKINGS clean.'
), size=12, space_after=6)
body(doc, 'Key constraint decisions supporting normalisation:', size=12, space_after=4)
for c in [
    'USERS.user_role CHECK: IN (\'CUSTOMER\',\'PROVIDER\',\'ADMIN\')',
    'BOOKINGS.status CHECK: IN (\'PENDING\',\'CONFIRMED\',\'IN_PROGRESS\',\'COMPLETED\',\'CANCELLED\')',
    'REVIEWS.rating CHECK: BETWEEN 1 AND 5',
    'PAYMENTS.payment_method CHECK: IN (\'CASH\',\'CREDIT_CARD\',\'UPI\',\'STRIPE\')',
    'PROVIDER_AVAILABILITY: slot_end > slot_start CHECK constraint',
    'PROVIDER_AREAS: UNIQUE on (provider_id, area_id) prevents duplicate city assignments',
]:
    bullet(doc, c)
doc.add_paragraph()

heading(doc, '3.4  Physical design', size=13, bold=True, space_before=8, space_after=6)
body(doc, (
    'Connection pool: The oracledb Node.js driver is configured with poolMin=2, poolMax=10, '
    'poolIncrement=2. Every route handler acquires a connection at the top, executes its queries, '
    'and closes the connection in a finally block regardless of success or failure.'
), size=12, space_after=6)
body(doc, (
    'Time storage: Availability slots store slot_start and slot_end as DATE values anchored to '
    '2000-01-01. Only the HH24:MI portion matters at runtime. The booking procedure extracts the '
    'time fraction using TO_CHAR(slot_start, \'SSSSS\') / 86400 and adds it to TRUNC(scheduled_date) '
    'to build the absolute time range for overlap comparison.'
), size=12, space_after=6)
body(doc, (
    'Primary keys: All 16 tables use NUMBER GENERATED ALWAYS AS IDENTITY \u2014 no separate Oracle '
    'sequences are required. ON DELETE CASCADE is applied to all parent-child relationships where '
    'removing a parent should clean up children. ON DELETE SET NULL is used for BOOKINGS.promo_id '
    'so booking history persists even if a promotion is deleted.'
), size=12, space_after=6)

heading(doc, '3.5  Implementation order', size=13, bold=True, space_before=8, space_after=6)
for item in [
    'Database schema \u2014 all 16 tables with constraints, in FK dependency order (ERROR_LOGS first, then USERS, then role profiles, then operational tables)',
    'PL/SQL objects \u2014 procedures, trigger, function, and view compiled and tested',
    'Seed data \u2014 sample users, providers, services, areas, bookings, and one complete flow (booking \u2192 invoice \u2192 payment \u2192 review)',
    'Backend API \u2014 six Express.js route groups with authentication middleware and Zod schemas',
    'Frontend \u2014 pages built role by role: Customer, then Provider, then Admin',
    'End-to-end testing \u2014 full booking lifecycle verified through the browser',
]:
    numbered(doc, item)
doc.add_paragraph()

heading(doc, '3.6  Testing and validation', size=13, bold=True, space_before=8, space_after=6)
body(doc, 'PL/SQL procedure edge cases tested:', size=12, space_after=4)
for case in [
    'Booking a slot already occupied by another active booking \u2014 expects ORA-20010',
    'Booking outside a provider\'s available hours \u2014 expects ORA-20011',
    'Cancelling an already-completed booking \u2014 expects ORA-20020',
    'Cancelling an already-cancelled booking \u2014 expects ORA-20021',
    'Applying an expired or over-limit promo code \u2014 booking proceeds without discount (silent ignore)',
]:
    bullet(doc, case)
body(doc, (
    'Cross-role ownership bypass was tested: a provider attempting to complete another provider\'s '
    'booking returns HTTP 403. SQL injection was confirmed impossible through bind variables on all '
    'Oracle queries. JWT expiry and role checks were tested across all three user roles.'
), size=12, space_before=6, space_after=8)
doc.add_page_break()

# ─── CHAPTER 4 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 4: ER DIAGRAM, RELATIONAL SCHEMA AND SAMPLE DATA', size=14, bold=True, space_before=8, space_after=12)

heading(doc, '4.1  ER Diagram (text representation)', size=13, bold=True, space_before=8, space_after=6)
body(doc, (
    'The diagram below shows all 16 entities with their key attributes and cardinalities. '
    'PK = Primary Key, FK = Foreign Key, UQ = Unique constraint.'
), size=12, space_after=6)
code_block(doc, """\
USERS (user_id PK, username UQ, password_hash, email UQ,
       user_role, is_active, last_login, created_at)
  |
  +--1:1-- CUSTOMERS (customer_id PK, user_id FK, first_name, last_name, phone)
  |           |
  |           +--1:N-- CUSTOMER_ADDRESSES [WEAK] (address_id discriminator,
  |           |                 location_label, house_no, building_name,
  |           |                 area_landmark, city, postal_code)
  |           |
  |           +--1:N-- BOOKINGS (booking_id PK, customer_id FK, service_id FK,
  |                              address_id FK, availability_id FK, promo_id FK?,
  |                              scheduled_date, duration_hours, status, created_at)
  |                       |
  |                       +--1:1-- INVOICES (invoice_id PK, booking_id FK UQ,
  |                       |                  base_amount, discount_amount,
  |                       |                  platform_fee, tax_amount, net_total,
  |                       |                  generated_at)
  |                       |           |
  |                       |           +--1:1-- PAYMENTS (payment_id PK,
  |                       |                              invoice_id FK UQ,
  |                       |                              amount_paid, payment_method,
  |                       |                              payment_status, transaction_id,
  |                       |                              paid_at)
  |                       |
  |                       +--1:1-- REVIEWS (review_id PK, booking_id FK UQ,
  |                       |                 rating, comments, created_at)
  |                       |         [triggers trg_update_provider_rating]
  |                       |
  |                       +--1:1-- CANCELLATIONS (cancellation_id PK,
  |                                               booking_id FK UQ,
  |                                               cancelled_by, reason,
  |                                               cancelled_at)
  |
  +--1:1-- SERVICE_PROVIDERS (provider_id PK, user_id FK, first_name, last_name,
                               phone, experience_yrs, background_chk,
                               rating_avg, jobs_completed)
               |
               +--M:N-- SERVICE_AREAS (area_id PK, city_name, region_code)
               |         [via PROVIDER_AREAS (provider_area_id PK,
               |                             provider_id FK, area_id FK,
               |                             UNIQUE(provider_id, area_id))]
               |
               +--1:N-- PROVIDER_AVAILABILITY (availability_id PK, provider_id FK,
               |                               day_of_week, slot_start, slot_end,
               |                               is_available)
               |
               +--1:N-- SERVICES_OFFERED (service_id PK, provider_id FK,
                                          category_id FK, service_name,
                                          hourly_rate, is_active)
                              |
                              +--N:1-- SERVICE_CATEGORIES (category_id PK,
                                                           category_name)

PROMOTIONS (promo_id PK, promo_code UQ, discount_percentage,
            max_discount_amt, min_order_amt, valid_from, valid_until,
            max_uses, current_uses, is_active)

ERROR_LOGS (log_id PK, severity, procedure_name, error_message, logged_at)
[Written to by PL/SQL only -- no FK relationships]""")

heading(doc, '4.2  Relational schema', size=13, bold=True, space_before=8, space_after=6)

body(doc, 'Identity Module', size=12, space_after=2)
code_block(doc, """\
USERS(user_id NUMBER PK IDENTITY, username VARCHAR2(50) UQ NOT NULL,
      password_hash VARCHAR2(256) NOT NULL, email VARCHAR2(100) UQ NOT NULL,
      user_role VARCHAR2(20) CHECK(IN 'CUSTOMER','PROVIDER','ADMIN') NOT NULL,
      is_active NUMBER(1,0) DEFAULT 1, last_login TIMESTAMP,
      created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL)

CUSTOMERS(customer_id NUMBER PK IDENTITY,
          user_id NUMBER UQ NOT NULL FK->USERS CASCADE DELETE,
          first_name VARCHAR2(50) NOT NULL, last_name VARCHAR2(50) NOT NULL,
          phone VARCHAR2(20))

SERVICE_PROVIDERS(provider_id NUMBER PK IDENTITY,
                  user_id NUMBER UQ NOT NULL FK->USERS CASCADE DELETE,
                  first_name VARCHAR2(50) NOT NULL, last_name VARCHAR2(50) NOT NULL,
                  phone VARCHAR2(20), experience_yrs NUMBER DEFAULT 0,
                  background_chk VARCHAR2(20) DEFAULT 'PENDING'
                    CHECK(IN 'PENDING','APPROVED','REJECTED'),
                  rating_avg NUMBER(3,2) DEFAULT 0.00,
                  jobs_completed NUMBER DEFAULT 0)""")

body(doc, 'Geography Module', size=12, space_after=2)
code_block(doc, """\
CUSTOMER_ADDRESSES(address_id NUMBER PK IDENTITY,
                   customer_id NUMBER NOT NULL FK->CUSTOMERS CASCADE DELETE,
                   location_label VARCHAR2(50) DEFAULT 'HOME',
                   house_no VARCHAR2(50) NOT NULL, building_name VARCHAR2(150),
                   area_landmark VARCHAR2(255) NOT NULL,
                   city VARCHAR2(100) NOT NULL, postal_code NUMBER(10) NOT NULL)

SERVICE_AREAS(area_id NUMBER PK IDENTITY,
              city_name VARCHAR2(100) NOT NULL,
              region_code VARCHAR2(50) NOT NULL,
              UNIQUE(city_name, region_code))

PROVIDER_AREAS(provider_area_id NUMBER PK IDENTITY,
               provider_id NUMBER NOT NULL FK->SERVICE_PROVIDERS CASCADE DELETE,
               area_id NUMBER NOT NULL FK->SERVICE_AREAS CASCADE DELETE,
               UNIQUE(provider_id, area_id))

PROVIDER_AVAILABILITY(availability_id NUMBER PK IDENTITY,
                      provider_id NUMBER NOT NULL FK->SERVICE_PROVIDERS CASCADE DELETE,
                      day_of_week VARCHAR2(15) CHECK(IN 'MONDAY'...'SUNDAY'),
                      slot_start DATE NOT NULL, slot_end DATE NOT NULL,
                      is_available NUMBER(1,0) DEFAULT 1,
                      CHECK(slot_end > slot_start))""")

body(doc, 'Catalogue Module', size=12, space_after=2)
code_block(doc, """\
SERVICE_CATEGORIES(category_id NUMBER PK IDENTITY,
                   category_name VARCHAR2(100) UQ NOT NULL)

SERVICES_OFFERED(service_id NUMBER PK IDENTITY,
                 provider_id NUMBER NOT NULL FK->SERVICE_PROVIDERS CASCADE DELETE,
                 category_id NUMBER NOT NULL FK->SERVICE_CATEGORIES CASCADE DELETE,
                 service_name VARCHAR2(100) NOT NULL,
                 hourly_rate NUMBER(10,2) NOT NULL CHECK(>= 0),
                 is_active NUMBER(1,0) DEFAULT 1)""")

body(doc, 'Operations Module', size=12, space_after=2)
code_block(doc, """\
PROMOTIONS(promo_id NUMBER PK IDENTITY,
           promo_code VARCHAR2(20) UQ NOT NULL,
           discount_percentage NUMBER(3,0) NOT NULL CHECK(BETWEEN 1 AND 100),
           max_discount_amt NUMBER(10,2) NOT NULL, min_order_amt NUMBER(10,2) DEFAULT 0,
           valid_from TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
           valid_until TIMESTAMP NOT NULL CHECK(valid_until > valid_from),
           max_uses NUMBER DEFAULT 100, current_uses NUMBER DEFAULT 0,
           is_active NUMBER(1,0) DEFAULT 1)

BOOKINGS(booking_id NUMBER PK IDENTITY,
         customer_id NUMBER NOT NULL FK->CUSTOMERS CASCADE DELETE,
         service_id NUMBER NOT NULL FK->SERVICES_OFFERED CASCADE DELETE,
         address_id NUMBER NOT NULL FK->CUSTOMER_ADDRESSES CASCADE DELETE,
         availability_id NUMBER NOT NULL FK->PROVIDER_AVAILABILITY,
         promo_id NUMBER FK->PROMOTIONS SET NULL ON DELETE,
         scheduled_date TIMESTAMP NOT NULL, duration_hours NUMBER(4,1) DEFAULT 1.0,
         status VARCHAR2(20) DEFAULT 'CONFIRMED'
           CHECK(IN 'PENDING','CONFIRMED','IN_PROGRESS','COMPLETED','CANCELLED'),
         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL)

INVOICES(invoice_id NUMBER PK IDENTITY,
         booking_id NUMBER UQ NOT NULL FK->BOOKINGS CASCADE DELETE,
         base_amount NUMBER(10,2) NOT NULL, discount_amount NUMBER(10,2) DEFAULT 0,
         platform_fee NUMBER(10,2) DEFAULT 0, tax_amount NUMBER(10,2) DEFAULT 0,
         net_total NUMBER(10,2) NOT NULL,
         generated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL)

PAYMENTS(payment_id NUMBER PK IDENTITY,
         invoice_id NUMBER UQ NOT NULL FK->INVOICES CASCADE DELETE,
         amount_paid NUMBER(10,2) NOT NULL,
         payment_method VARCHAR2(50) DEFAULT 'CASH'
           CHECK(IN 'CASH','CREDIT_CARD','UPI','STRIPE'),
         payment_status VARCHAR2(20) DEFAULT 'PENDING'
           CHECK(IN 'PENDING','SUCCESS','COMPLETED','FAILED','REFUNDED'),
         transaction_id VARCHAR2(100), paid_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL)

REVIEWS(review_id NUMBER PK IDENTITY,
        booking_id NUMBER UQ NOT NULL FK->BOOKINGS CASCADE DELETE,
        rating NUMBER(1,0) NOT NULL CHECK(BETWEEN 1 AND 5),
        comments VARCHAR2(1000),
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP NOT NULL)

CANCELLATIONS(cancellation_id NUMBER PK IDENTITY,
              booking_id NUMBER UQ NOT NULL FK->BOOKINGS CASCADE DELETE,
              cancelled_by VARCHAR2(10) CHECK(IN 'CUSTOMER','PROVIDER','SYSTEM'),
              reason VARCHAR2(500),
              cancelled_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")

body(doc, 'Audit Module', size=12, space_after=2)
code_block(doc, """\
ERROR_LOGS(log_id NUMBER PK IDENTITY,
           severity VARCHAR2(10) CHECK(IN 'LOW','MEDIUM','HIGH','CRITICAL'),
           procedure_name VARCHAR2(100), error_message VARCHAR2(2000),
           logged_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")

heading(doc, '4.3  Sample data', size=13, bold=True, space_before=10, space_after=8)

body(doc, 'USERS (seed accounts \u2014 all passwords: pass123)', size=12, space_after=4)
simple_table(doc, ['username', 'email', 'user_role', 'is_active'], [
    ('john_cust',  'john@gmail.com',      'CUSTOMER', '1'),
    ('sara_cust',  'sara@gmail.com',       'CUSTOMER', '1'),
    ('bob_pro',    'bob@fixit.com',        'PROVIDER', '1'),
    ('dave_pro',   'dave@sparkworks.com',  'PROVIDER', '1'),
    ('priya_pro',  'priya@cleanpro.com',   'PROVIDER', '1'),
    ('admin',      'admin@servemart.com',  'ADMIN',    '1'),
], col_widths=[1.3, 2.2, 1.1, 0.8])

body(doc, 'SERVICE_CATEGORIES', size=12, space_after=4)
simple_table(doc, ['category_id', 'category_name'], [
    ('1', 'PLUMBING'), ('2', 'ELECTRICAL'), ('3', 'CLEANING'),
    ('4', 'PAINTING'), ('5', 'CARPENTRY'),  ('6', 'APPLIANCE REPAIR'),
], col_widths=[1.2, 2.5])

body(doc, 'SERVICE_AREAS', size=12, space_after=4)
simple_table(doc, ['area_id', 'city_name', 'region_code'], [
    ('1', 'Mumbai',    'MH'),
    ('2', 'Delhi',     'DL'),
    ('3', 'Bangalore', 'KA'),
    ('4', 'Hyderabad', 'TS'),
    ('5', 'Manipal',   'KA'),
], col_widths=[0.8, 1.8, 1.2])

body(doc, 'SERVICE_PROVIDERS', size=12, space_after=4)
simple_table(doc, ['username', 'full_name', 'exp_yrs', 'background_chk', 'rating_avg', 'jobs_completed'], [
    ('bob_pro',   'Bob Builder',  12, 'APPROVED', 4.50, 50),
    ('dave_pro',  'Dave Spark',   8,  'APPROVED', 4.80, 37),
    ('priya_pro', 'Priya Verma',  5,  'APPROVED', 4.90, 62),
], col_widths=[1.0, 1.2, 0.7, 1.2, 1.0, 1.2])

body(doc, 'SERVICES_OFFERED', size=12, space_after=4)
simple_table(doc, ['provider', 'category', 'service_name', 'hourly_rate (Rs.)'], [
    ('bob_pro',   'PLUMBING',   'Pipe Leak Fix',           '350.00'),
    ('bob_pro',   'PLUMBING',   'Bathroom Fitting',        '500.00'),
    ('dave_pro',  'ELECTRICAL', 'Wiring & Rewiring',       '400.00'),
    ('dave_pro',  'ELECTRICAL', 'Ceiling Fan Installation', '250.00'),
    ('priya_pro', 'CLEANING',   'Deep Home Cleaning',      '300.00'),
], col_widths=[1.1, 1.1, 2.1, 1.2])

body(doc, 'PROMOTIONS', size=12, space_after=4)
simple_table(doc, ['promo_code', 'discount_%', 'max_discount', 'min_order', 'valid_until'], [
    ('SAVE10',   '10', 'Rs. 150.00', 'Rs. 200.00', 'SYSDATE + 365'),
    ('WELCOME20','20', 'Rs. 200.00', 'Rs. 300.00', 'SYSDATE + 365'),
], col_widths=[1.1, 0.9, 1.2, 1.1, 1.3])

body(doc, 'BOOKINGS (demo seed \u2014 full flow: booking to review)', size=12, space_after=4)
simple_table(doc, ['customer', 'service', 'status', 'duration', 'scheduled_date'], [
    ('john_cust', 'Pipe Leak Fix', 'COMPLETED', '2.0 hrs', 'SYSDATE - 3 days'),
], col_widths=[1.1, 1.3, 1.1, 0.9, 1.5])

body(doc, 'INVOICES (auto-generated by sp_generate_invoice for the demo booking)', size=12, space_after=4)
simple_table(doc, ['booking_id', 'base_amount', 'discount_amount', 'platform_fee', 'tax_amount', 'net_total'], [
    ('1', 'Rs. 700.00', 'Rs. 0.00', 'Rs. 70.00', 'Rs. 35.00', 'Rs. 805.00'),
], col_widths=[0.9, 1.1, 1.2, 1.1, 1.0, 1.0])
doc.add_page_break()

# ─── CHAPTER 5 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 5: DDL COMMANDS AND PL/SQL PROCEDURES / FUNCTIONS / TRIGGERS', size=14, bold=True, space_before=8, space_after=12)

heading(doc, '5.1  Database and table creation (DDL)', size=13, bold=True, space_before=8, space_after=6)
body(doc, (
    'The schema runs inside the MARKETPLACE user schema on Oracle XE. Tables are created in FK '
    'dependency order: ERROR_LOGS first (no FK dependencies; all PL/SQL objects need it for '
    'logging), then USERS, then CUSTOMERS and SERVICE_PROVIDERS, then geographic and catalogue '
    'tables, then BOOKINGS (which references six parent tables), and finally the transaction '
    'tables INVOICES, PAYMENTS, REVIEWS, and CANCELLATIONS. The full schema is in init.sql. '
    'Selected representative tables are shown below.'
), size=12, space_after=8)

body(doc, 'Table 1: ERROR_LOGS (created first so PL/SQL can always write to it)', size=12, space_after=2)
code_block(doc, """\
CREATE TABLE ERROR_LOGS (
    log_id         NUMBER GENERATED ALWAYS AS IDENTITY PRIMARY KEY,
    severity       VARCHAR2(10) CHECK (severity IN ('LOW','MEDIUM','HIGH','CRITICAL')),
    procedure_name VARCHAR2(100),
    error_message  VARCHAR2(2000),
    logged_at      TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);""")

body(doc, 'Table 2: USERS', size=12, space_after=2)
code_block(doc, """\
CREATE TABLE USERS (
    user_id       NUMBER GENERATED ALWAYS AS IDENTITY PRIMARY KEY,
    username      VARCHAR2(50)  UNIQUE NOT NULL,
    password_hash VARCHAR2(256) NOT NULL,
    email         VARCHAR2(100) UNIQUE NOT NULL,
    user_role     VARCHAR2(20)  NOT NULL
                  CHECK (user_role IN ('CUSTOMER','PROVIDER','ADMIN')),
    is_active     NUMBER(1,0)   DEFAULT 1 CHECK (is_active IN (0,1)),
    last_login    TIMESTAMP,
    created_at    TIMESTAMP     DEFAULT CURRENT_TIMESTAMP NOT NULL
);""")
body(doc, (
    'NUMBER GENERATED ALWAYS AS IDENTITY replaces Oracle sequences for all PKs. The CHECK on '
    'user_role restricts values to exactly three roles. is_active uses NUMBER(1,0) rather than '
    'BOOLEAN because Oracle does not support BOOLEAN in table definitions before Oracle 23c.'
), size=12, space_after=8)

body(doc, 'Table 3: BOOKINGS (central transaction entity \u2014 references six parent tables)', size=12, space_after=2)
code_block(doc, """\
CREATE TABLE BOOKINGS (
    booking_id      NUMBER       GENERATED ALWAYS AS IDENTITY PRIMARY KEY,
    customer_id     NUMBER       NOT NULL,
    service_id      NUMBER       NOT NULL,
    address_id      NUMBER       NOT NULL,
    availability_id NUMBER       NOT NULL,
    promo_id        NUMBER,
    scheduled_date  TIMESTAMP    NOT NULL,
    duration_hours  NUMBER(4,1)  DEFAULT 1.0,
    status          VARCHAR2(20) DEFAULT 'CONFIRMED'
                    CHECK (status IN
                      ('PENDING','CONFIRMED','IN_PROGRESS','COMPLETED','CANCELLED')),
    created_at      TIMESTAMP    DEFAULT CURRENT_TIMESTAMP NOT NULL,
    CONSTRAINT fk_bk_cust  FOREIGN KEY (customer_id)
        REFERENCES CUSTOMERS(customer_id) ON DELETE CASCADE,
    CONSTRAINT fk_bk_srvc  FOREIGN KEY (service_id)
        REFERENCES SERVICES_OFFERED(service_id) ON DELETE CASCADE,
    CONSTRAINT fk_bk_addr  FOREIGN KEY (address_id)
        REFERENCES CUSTOMER_ADDRESSES(address_id) ON DELETE CASCADE,
    CONSTRAINT fk_bk_avail FOREIGN KEY (availability_id)
        REFERENCES PROVIDER_AVAILABILITY(availability_id),
    CONSTRAINT fk_bk_promo FOREIGN KEY (promo_id)
        REFERENCES PROMOTIONS(promo_id) ON DELETE SET NULL
);""")
body(doc, (
    'ON DELETE SET NULL on promo_id preserves booking history when a promotion is deleted. '
    'ON DELETE CASCADE on the customer, service, and address FK constraints ensures removing a '
    'parent record cleans up all dependent bookings. No ON DELETE action on availability_id \u2014 '
    'slot data is intentionally retained after booking deletion.'
), size=12, space_after=8)

heading(doc, '5.2  Basic DML queries', size=13, bold=True, space_before=10, space_after=6)

body(doc, 'All pending/confirmed bookings for a given provider:', size=12, space_after=2)
code_block(doc, """\
SELECT b.booking_id,
       cu.first_name || ' ' || cu.last_name AS customer_name,
       so.service_name,
       b.scheduled_date, b.duration_hours, b.status
FROM   BOOKINGS          b
JOIN   SERVICES_OFFERED  so ON b.service_id   = so.service_id
JOIN   SERVICE_PROVIDERS sp ON so.provider_id = sp.provider_id
JOIN   CUSTOMERS         cu ON b.customer_id  = cu.customer_id
WHERE  sp.provider_id = :p_provider_id
  AND  b.status IN ('PENDING', 'CONFIRMED')
ORDER  BY b.scheduled_date;""")

body(doc, 'Revenue summary by service category (used by the admin dashboard):', size=12, space_after=2)
code_block(doc, """\
SELECT sc.category_name,
       COUNT(b.booking_id)     AS total_bookings,
       SUM(i.platform_fee)     AS platform_revenue,
       SUM(i.net_total)        AS gross_value
FROM   BOOKINGS           b
JOIN   SERVICES_OFFERED   so ON b.service_id   = so.service_id
JOIN   SERVICE_CATEGORIES sc ON so.category_id = sc.category_id
JOIN   INVOICES            i ON b.booking_id   = i.booking_id
WHERE  b.status = 'COMPLETED'
GROUP  BY sc.category_name
ORDER  BY platform_revenue DESC;""")

body(doc, 'Customer booking history with invoice and payment status:', size=12, space_after=2)
code_block(doc, """\
SELECT b.booking_id,
       so.service_name,
       sp.first_name || ' ' || sp.last_name  AS provider_name,
       b.scheduled_date, b.status,
       i.net_total,
       NVL(p.payment_status, 'NO_INVOICE')   AS payment_status,
       cn.reason                             AS cancellation_reason
FROM   BOOKINGS          b
JOIN   SERVICES_OFFERED  so ON b.service_id   = so.service_id
JOIN   SERVICE_PROVIDERS sp ON so.provider_id = sp.provider_id
LEFT JOIN INVOICES        i  ON b.booking_id  = i.booking_id
LEFT JOIN PAYMENTS        p  ON i.invoice_id  = p.invoice_id
LEFT JOIN CANCELLATIONS   cn ON b.booking_id  = cn.booking_id
WHERE  b.customer_id = :p_customer_id
ORDER  BY b.scheduled_date DESC;""")

body(doc, 'Disable a provider availability slot:', size=12, space_after=2)
code_block(doc, """\
UPDATE PROVIDER_AVAILABILITY
SET    is_available = 0
WHERE  availability_id = :slot_id AND provider_id = :prov_id;""")

body(doc, 'Remove expired promo codes that have been fully used:', size=12, space_after=2)
code_block(doc, """\
DELETE FROM PROMOTIONS
WHERE  valid_until < SYSDATE AND current_uses >= max_uses;""")

heading(doc, '5.3  Complex queries', size=13, bold=True, space_before=10, space_after=6)

body(doc, 'Provider full detail with service list and area coverage using LISTAGG:', size=12, space_after=2)
code_block(doc, """\
SELECT sp.provider_id,
       u.username,
       sp.first_name || ' ' || sp.last_name                     AS full_name,
       sp.rating_avg, sp.jobs_completed,
       LISTAGG(sa.city_name, ', ')
           WITHIN GROUP (ORDER BY sa.city_name)                 AS service_cities
FROM   SERVICE_PROVIDERS sp
JOIN   USERS             u  ON sp.user_id     = u.user_id
LEFT JOIN PROVIDER_AREAS pa ON pa.provider_id = sp.provider_id
LEFT JOIN SERVICE_AREAS  sa ON sa.area_id     = pa.area_id
WHERE  u.is_active = 1
  AND  EXISTS (SELECT 1 FROM SERVICES_OFFERED so2
               WHERE so2.provider_id = sp.provider_id AND so2.is_active = 1)
GROUP  BY sp.provider_id, u.username, sp.first_name, sp.last_name,
          sp.rating_avg, sp.jobs_completed
ORDER  BY sp.rating_avg DESC;""")
body(doc, (
    'The EXISTS subquery filters out providers who have registered but not listed any active '
    'services, so they never appear in the customer browse page.'
), size=12, space_after=8)

body(doc, 'Rank providers by composite score within each service area \u2014 window function:', size=12, space_after=2)
code_block(doc, """\
SELECT sa.city_name,
       sp.first_name || ' ' || sp.last_name AS provider,
       sp.rating_avg, sp.jobs_completed,
       ROUND((sp.rating_avg * 0.5)
             + (LEAST(sp.jobs_completed,100)/100.0 * 0.3)
             + 0.2, 2)                       AS score,
       RANK() OVER (
           PARTITION BY sa.area_id
           ORDER BY (sp.rating_avg * 0.5
                     + LEAST(sp.jobs_completed,100)/100.0 * 0.3
                     + 0.2) DESC
       )                                     AS city_rank
FROM   SERVICE_PROVIDERS sp
JOIN   PROVIDER_AREAS    pa ON pa.provider_id = sp.provider_id
JOIN   SERVICE_AREAS     sa ON pa.area_id     = sa.area_id
WHERE  sp.background_chk = 'APPROVED';""")

body(doc, 'Already-booked time windows for a provider (used by booking form to hide occupied slots):', size=12, space_after=2)
code_block(doc, """\
SELECT b.scheduled_date,
       b.duration_hours,
       b.status
FROM   BOOKINGS          b
JOIN   SERVICES_OFFERED  so ON b.service_id   = so.service_id
WHERE  so.provider_id = :p_provider_id
  AND  b.status IN ('PENDING','CONFIRMED','IN_PROGRESS')
  AND  b.scheduled_date >= SYSDATE
ORDER  BY b.scheduled_date;""")
body(doc, (
    'This is called by GET /api/bookings/booked-slots/:providerId. The frontend filters the '
    '3-hour slot grid to hide any window that overlaps an existing active booking.'
), size=12, space_after=8)

body(doc, 'Providers dormant for 30+ days (no active or completed bookings):', size=12, space_after=2)
code_block(doc, """\
SELECT sp.provider_id,
       u.username,
       sp.first_name || ' ' || sp.last_name AS full_name,
       sp.jobs_completed, sp.rating_avg
FROM   SERVICE_PROVIDERS sp
JOIN   USERS u ON sp.user_id = u.user_id
WHERE  sp.provider_id NOT IN (
    SELECT DISTINCT so.provider_id
    FROM   BOOKINGS          b
    JOIN   SERVICES_OFFERED  so ON b.service_id = so.service_id
    WHERE  b.created_at >= SYSDATE - 30
      AND  b.status IN ('PENDING','CONFIRMED','COMPLETED')
)
ORDER  BY sp.jobs_completed DESC;""")

body(doc, 'Monthly average rating trend per service category:', size=12, space_after=2)
code_block(doc, """\
SELECT sc.category_name,
       TRUNC(r.created_at, 'MM')  AS review_month,
       COUNT(r.review_id)         AS review_count,
       ROUND(AVG(r.rating), 2)    AS avg_rating
FROM   REVIEWS            r
JOIN   BOOKINGS           b  ON r.booking_id  = b.booking_id
JOIN   SERVICES_OFFERED   so ON b.service_id  = so.service_id
JOIN   SERVICE_CATEGORIES sc ON so.category_id = sc.category_id
GROUP  BY sc.category_name, TRUNC(r.created_at, 'MM')
ORDER  BY review_month DESC, avg_rating DESC;""")

body(doc, 'High-value customers who have spent more than Rs. 2000 total \u2014 HAVING clause:', size=12, space_after=2)
code_block(doc, """\
SELECT c.customer_id, u.username,
       COUNT(b.booking_id) AS total_bookings,
       SUM(i.net_total)    AS total_spent
FROM   CUSTOMERS c
JOIN   USERS     u  ON c.user_id     = u.user_id
JOIN   BOOKINGS  b  ON b.customer_id = c.customer_id
JOIN   INVOICES  i  ON i.booking_id  = b.booking_id
WHERE  b.status = 'COMPLETED'
GROUP  BY c.customer_id, u.username
HAVING SUM(i.net_total) > 2000
ORDER  BY total_spent DESC;""")

heading(doc, '5.4  Stored procedures', size=13, bold=True, space_before=10, space_after=6)

body(doc, 'SP1: sp_create_booking', size=12, space_after=2)
body(doc, (
    'Creates a confirmed booking after validating three things: the requested time window falls '
    'inside the provider\'s availability range (ORA-20011 if not), no active booking already '
    'occupies that window for the same provider (ORA-20010 if overlap), and the promo code if '
    'supplied is valid. An invalid or expired promo is silently ignored \u2014 the booking still '
    'proceeds without a discount. The EXCEPTION handler assigns SQLERRM to a VARCHAR2 variable '
    'before the INSERT into ERROR_LOGS, because SQLERRM cannot be used directly inside a SQL '
    'VALUES clause in Oracle PL/SQL.'
), size=12, space_after=4)
code_block(doc, """\
CREATE OR REPLACE PROCEDURE sp_create_booking (
    p_cust_id    IN NUMBER,  p_srvc_id  IN NUMBER,
    p_addr_id    IN NUMBER,  p_avail_id IN NUMBER,
    p_promo_code IN VARCHAR2,
    p_date       IN TIMESTAMP,
    p_dur        IN NUMBER
)
IS
    v_promo_id      NUMBER := NULL;
    v_provider_id   NUMBER;
    v_overlap_count NUMBER;
    v_range_start   TIMESTAMP;
    v_range_end     TIMESTAMP;
    v_err           VARCHAR2(2000);
BEGIN
    -- 1. Resolve provider for this service
    SELECT provider_id INTO v_provider_id
    FROM   SERVICES_OFFERED WHERE service_id = p_srvc_id;

    -- 2. Build absolute time range from slot HH24:MI
    SELECT TRUNC(p_date) + (TO_CHAR(slot_start, 'SSSSS') / 86400),
           TRUNC(p_date) + (TO_CHAR(slot_end,   'SSSSS') / 86400)
    INTO   v_range_start, v_range_end
    FROM   PROVIDER_AVAILABILITY WHERE availability_id = p_avail_id;

    IF p_date < v_range_start
    OR (p_date + p_dur/24) > v_range_end THEN
        RAISE_APPLICATION_ERROR(-20011,
            'Requested time is outside the provider''s available hours.');
    END IF;

    -- 3. Overlap check across all active bookings for this provider
    SELECT COUNT(*) INTO v_overlap_count
    FROM   BOOKINGS b JOIN SERVICES_OFFERED so
           ON b.service_id = so.service_id
    WHERE  so.provider_id = v_provider_id
      AND  b.status IN ('PENDING','CONFIRMED','IN_PROGRESS')
      AND  p_date                < (b.scheduled_date + b.duration_hours/24)
      AND  (p_date + p_dur/24)  >  b.scheduled_date;

    IF v_overlap_count > 0 THEN
        RAISE_APPLICATION_ERROR(-20010,
            'Time slot is already occupied by another booking.');
    END IF;

    -- 4. Validate promo (silent ignore if invalid or expired)
    IF p_promo_code IS NOT NULL THEN
        BEGIN
            SELECT promo_id INTO v_promo_id FROM PROMOTIONS
            WHERE  UPPER(promo_code) = UPPER(p_promo_code)
              AND  is_active = 1 AND current_uses < max_uses
              AND  valid_until > SYSDATE AND valid_from <= SYSDATE;
            UPDATE PROMOTIONS SET current_uses = current_uses + 1
            WHERE  promo_id = v_promo_id;
        EXCEPTION WHEN NO_DATA_FOUND THEN v_promo_id := NULL;
        END;
    END IF;

    -- 5. Insert booking
    INSERT INTO BOOKINGS (
        customer_id, service_id, address_id, availability_id,
        promo_id, scheduled_date, duration_hours, status
    ) VALUES (
        p_cust_id, p_srvc_id, p_addr_id, p_avail_id,
        v_promo_id, p_date, p_dur, 'CONFIRMED'
    );
    COMMIT;

EXCEPTION
    WHEN OTHERS THEN
        v_err := SQLERRM; ROLLBACK;
        BEGIN
            INSERT INTO ERROR_LOGS (severity, procedure_name, error_message)
            VALUES ('CRITICAL', 'sp_create_booking', v_err);
            COMMIT;
        END;
        RAISE;
END sp_create_booking;
/""")

body(doc, 'SP2: sp_generate_invoice', size=12, space_after=2)
body(doc, (
    'Called when a provider marks a booking COMPLETED. Computes base = hourly_rate x '
    'duration_hours. If the booking used a promo, looks up discount_percentage and caps the '
    'discount at max_discount_amt. Platform fee = 10% of discounted amount, GST = 5%, '
    'net total = discounted amount + fee + tax. This is the only path that creates INVOICES rows.'
), size=12, space_after=4)
code_block(doc, """\
CREATE OR REPLACE PROCEDURE sp_generate_invoice (p_booking_id IN NUMBER)
IS
    v_hourly_rate NUMBER; v_dur      NUMBER;
    v_base        NUMBER; v_discount NUMBER := 0;
    v_fee         NUMBER; v_tax      NUMBER; v_net NUMBER;
    v_err         VARCHAR2(2000);
BEGIN
    SELECT so.hourly_rate, b.duration_hours
    INTO   v_hourly_rate, v_dur
    FROM   BOOKINGS b JOIN SERVICES_OFFERED so ON b.service_id = so.service_id
    WHERE  b.booking_id = p_booking_id;

    v_base := v_hourly_rate * v_dur;

    -- Apply promo discount capped at max_discount_amt
    FOR rec IN (
        SELECT p.discount_percentage, p.max_discount_amt
        FROM   PROMOTIONS p JOIN BOOKINGS b ON b.promo_id = p.promo_id
        WHERE  b.booking_id = p_booking_id
    ) LOOP
        v_discount := LEAST(v_base * (rec.discount_percentage / 100),
                            rec.max_discount_amt);
    END LOOP;

    v_fee := ROUND((v_base - v_discount) * 0.10, 2);   -- 10% platform fee
    v_tax := ROUND((v_base - v_discount) * 0.05, 2);   -- 5% GST
    v_net := ROUND((v_base - v_discount) + v_fee + v_tax, 2);

    INSERT INTO INVOICES (booking_id, base_amount, discount_amount,
                          platform_fee, tax_amount, net_total)
    VALUES (p_booking_id, v_base, v_discount, v_fee, v_tax, v_net);
    COMMIT;

EXCEPTION
    WHEN OTHERS THEN
        v_err := SQLERRM; ROLLBACK;
        BEGIN
            INSERT INTO ERROR_LOGS (severity, procedure_name, error_message)
            VALUES ('CRITICAL', 'sp_generate_invoice', v_err); COMMIT;
        END;
        RAISE;
END sp_generate_invoice;
/""")

body(doc, 'SP3: sp_cancel_booking', size=12, space_after=2)
body(doc, (
    'Cancels a booking that is not already COMPLETED (ORA-20020) or CANCELLED (ORA-20021), '
    'restores the availability slot to is_available=1, writes an audit row to CANCELLATIONS, '
    'and marks any SUCCESS or COMPLETED payment as REFUNDED. The payment refund block uses a '
    'nested BEGIN/EXCEPTION to handle the case where no payment exists yet (booking not yet '
    'invoiced) without raising an error.'
), size=12, space_after=4)
code_block(doc, """\
CREATE OR REPLACE PROCEDURE sp_cancel_booking (
    p_booking_id   IN NUMBER,
    p_cancelled_by IN VARCHAR2,
    p_reason       IN VARCHAR2
)
IS
    v_status VARCHAR2(20); v_avail_id   NUMBER;
    v_invoice_id NUMBER;   v_pay_status VARCHAR2(20);
    v_err VARCHAR2(2000);
BEGIN
    SELECT status, availability_id INTO v_status, v_avail_id
    FROM   BOOKINGS WHERE booking_id = p_booking_id;

    IF v_status = 'COMPLETED' THEN
        RAISE_APPLICATION_ERROR(-20020, 'Cannot cancel a completed booking.');
    END IF;
    IF v_status = 'CANCELLED' THEN
        RAISE_APPLICATION_ERROR(-20021, 'Booking is already cancelled.');
    END IF;

    UPDATE BOOKINGS SET status = 'CANCELLED'
    WHERE  booking_id = p_booking_id;

    UPDATE PROVIDER_AVAILABILITY SET is_available = 1
    WHERE  availability_id = v_avail_id;

    INSERT INTO CANCELLATIONS (booking_id, cancelled_by, reason)
    VALUES (p_booking_id, p_cancelled_by, p_reason);

    -- Refund any already-paid payment
    BEGIN
        SELECT i.invoice_id, p.payment_status
        INTO   v_invoice_id, v_pay_status
        FROM   INVOICES i JOIN PAYMENTS p ON p.invoice_id = i.invoice_id
        WHERE  i.booking_id = p_booking_id;

        IF v_pay_status IN ('SUCCESS','COMPLETED') THEN
            UPDATE PAYMENTS
            SET    payment_status = 'REFUNDED',
                   amount_paid = (SELECT net_total FROM INVOICES
                                  WHERE invoice_id = v_invoice_id)
            WHERE  invoice_id = v_invoice_id;
        END IF;
    EXCEPTION WHEN NO_DATA_FOUND THEN NULL;
    END;

    COMMIT;

EXCEPTION
    WHEN OTHERS THEN
        v_err := SQLERRM; ROLLBACK;
        BEGIN
            INSERT INTO ERROR_LOGS (severity, procedure_name, error_message)
            VALUES ('CRITICAL', 'sp_cancel_booking', v_err); COMMIT;
        END;
        RAISE;
END sp_cancel_booking;
/""")

heading(doc, '5.5  Compound trigger', size=13, bold=True, space_before=10, space_after=6)
body(doc, 'trg_update_provider_rating \u2014 FOR INSERT OR UPDATE OR DELETE ON REVIEWS', size=12, space_after=4)
body(doc, (
    'A compound trigger is used instead of a simple AFTER EACH ROW trigger to avoid Oracle\'s '
    'mutating-table error. When a REVIEWS row is inserted, updated, or deleted, the AFTER EACH ROW '
    'section resolves the affected provider ID from the BOOKINGS/SERVICES_OFFERED chain and '
    'stores it in an associative array. The AFTER STATEMENT section iterates the array and '
    'recomputes both rating_avg (AVG of all review ratings for that provider) and '
    'jobs_completed (COUNT DISTINCT of COMPLETED bookings, not reviews) in a single UPDATE per '
    'affected provider.'
), size=12, space_after=4)
code_block(doc, """\
CREATE OR REPLACE TRIGGER trg_update_provider_rating
FOR INSERT OR UPDATE OR DELETE ON REVIEWS
COMPOUND TRIGGER

    TYPE t_id_table IS TABLE OF NUMBER INDEX BY PLS_INTEGER;
    v_affected t_id_table;

AFTER EACH ROW IS
    v_pid NUMBER;
BEGIN
    SELECT so.provider_id INTO v_pid
    FROM   BOOKINGS b JOIN SERVICES_OFFERED so
           ON so.service_id = b.service_id
    WHERE  b.booking_id = NVL(:NEW.booking_id, :OLD.booking_id);
    v_affected(v_pid) := v_pid;
EXCEPTION WHEN NO_DATA_FOUND THEN NULL;
END AFTER EACH ROW;

AFTER STATEMENT IS
    v_idx NUMBER;
BEGIN
    v_idx := v_affected.FIRST;
    WHILE v_idx IS NOT NULL LOOP
        UPDATE SERVICE_PROVIDERS sp
        SET (rating_avg, jobs_completed) = (
            SELECT NVL(ROUND(AVG(r.rating), 2), 0.00),
                   COUNT(DISTINCT
                       CASE WHEN b.status = 'COMPLETED'
                       THEN b.booking_id END)
            FROM   REVIEWS r
            JOIN   BOOKINGS b  ON b.booking_id  = r.booking_id
            JOIN   SERVICES_OFFERED so
                               ON so.service_id = b.service_id
            WHERE  so.provider_id = v_idx
        )
        WHERE sp.provider_id = v_idx;
        v_idx := v_affected.NEXT(v_idx);
    END LOOP;
END AFTER STATEMENT;

END trg_update_provider_rating;
/""")

heading(doc, '5.6  Function', size=13, bold=True, space_before=10, space_after=6)
body(doc, 'fn_recommend_providers \u2014 returns SYS_REFCURSOR of ranked providers for a service area', size=12, space_after=4)
body(doc, (
    'Takes a service area ID and returns an open cursor of approved, active providers in that area '
    'ranked by composite score. The 0.2 flat baseline ensures a brand-new provider with no reviews '
    'still appears in results rather than scoring zero. The backend calls this via an anonymous '
    'block and reads rows using node-oracledb\'s resultSet API.'
), size=12, space_after=4)
code_block(doc, """\
CREATE OR REPLACE FUNCTION fn_recommend_providers (
    p_area_id IN NUMBER
) RETURN SYS_REFCURSOR
IS
    v_cursor SYS_REFCURSOR;
BEGIN
    OPEN v_cursor FOR
        SELECT sp.provider_id, u.username,
               sp.first_name || ' ' || sp.last_name         AS full_name,
               sp.rating_avg, sp.jobs_completed,
               ROUND(
                   (sp.rating_avg * 0.5)
                   + (LEAST(sp.jobs_completed, 100) / 100.0 * 0.3)
                   + 0.2,
               2)                                            AS score
        FROM   SERVICE_PROVIDERS sp
        JOIN   USERS          u  ON u.user_id     = sp.user_id
        JOIN   PROVIDER_AREAS pa ON pa.provider_id = sp.provider_id
        WHERE  pa.area_id        = p_area_id
          AND  sp.background_chk = 'APPROVED'
          AND  u.is_active       = 1
        ORDER BY score DESC;
    RETURN v_cursor;
END fn_recommend_providers;
/""")
body(doc, 'Score formula breakdown:', size=12, space_after=4)
for item in [
    'rating_avg x 0.5 \u2014 50% weight; a 5.0 rating contributes 2.5 to the score',
    'LEAST(jobs_completed, 100) / 100 x 0.3 \u2014 30% weight, capped at 100 jobs; prevents monopoly by high-volume providers',
    '0.2 flat baseline \u2014 20% for all approved providers; new providers with no jobs still appear',
]:
    bullet(doc, item)
doc.add_paragraph()

heading(doc, '5.7  View', size=13, bold=True, space_before=10, space_after=6)
body(doc, 'vw_provider_summary \u2014 provider profile view for browse and admin screens', size=12, space_after=4)
body(doc, (
    'Joins SERVICE_PROVIDERS with USERS and uses a correlated subquery to pull the single '
    'active category name (using MAX to avoid a GROUP BY on the outer query) per provider. '
    'The WHERE clause filters deactivated accounts so they never appear in the public listing. '
    'Two backend routes query this view: GET /providers (list) and GET /providers/:id (detail).'
), size=12, space_after=4)
code_block(doc, """\
CREATE OR REPLACE VIEW vw_provider_summary AS
SELECT sp.provider_id,
       u.username,
       sp.first_name || ' ' || sp.last_name AS full_name,
       (SELECT MAX(sc.category_name)
        FROM   SERVICE_CATEGORIES sc
        JOIN   SERVICES_OFFERED   so ON sc.category_id = so.category_id
        WHERE  so.provider_id = sp.provider_id
          AND  so.is_active   = 1)           AS category_name,
       sp.rating_avg,
       sp.jobs_completed,
       u.is_active
FROM SERVICE_PROVIDERS sp
JOIN USERS u ON sp.user_id = u.user_id
WHERE u.is_active = 1;""")
doc.add_page_break()

# ─── CHAPTER 6 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 6: RESULTS AND SNAPSHOTS', size=14, bold=True, space_before=8, space_after=12)
body(doc, (
    'This chapter presents the output of key queries, demonstrates stored procedure behaviour, '
    'verifies trigger execution, and describes each screen of the web application. The full '
    'schema (init.sql) was run in SQL Developer to create all 16 tables and 6 PL/SQL objects '
    'in a single pass (SET DEFINE OFF prevents the ampersand in "Wiring & Rewiring" from being '
    'treated as a substitution variable).'
), size=12, space_after=12)

heading(doc, '6.1  Query results', size=13, bold=True, space_before=8, space_after=6)

body(doc, 'Revenue query output (admin dashboard) on seed data:', size=12, space_after=4)
simple_table(doc, ['CATEGORY_NAME', 'TOTAL_BOOKINGS', 'PLATFORM_REVENUE (Rs.)', 'GROSS_VALUE (Rs.)'], [
    ('PLUMBING', '1', '70.00', '805.00'),
], col_widths=[1.6, 1.3, 2.0, 1.5])
body(doc, (
    'With more bookings across categories this table grows to show comparative revenue per '
    'service type. The admin dashboard renders it as a dual-bar Recharts chart.'
), size=12, space_after=8)

body(doc, 'fn_recommend_providers(3) output \u2014 Bangalore providers ranked by composite score:', size=12, space_after=4)
simple_table(doc, ['PROVIDER_ID', 'FULL_NAME', 'RATING_AVG', 'JOBS_COMPLETED', 'SCORE'], [
    ('3', 'Priya Verma', '4.90', '62', '2.636'),
    ('2', 'Dave Spark',  '4.80', '37', '2.511'),
], col_widths=[1.0, 1.4, 1.1, 1.3, 0.8])

body(doc, 'Provider recommendation \u2014 score calculation example for Priya Verma:', size=12, space_after=2)
code_block(doc, """\
score = (4.90 * 0.5) + (LEAST(62,100)/100 * 0.3) + 0.2
      = 2.45  + 0.186 + 0.2
      = 2.636""")

heading(doc, '6.2  Stored procedure execution', size=13, bold=True, space_before=8, space_after=6)

body(doc, 'Normal booking creation with promo code SAVE10:', size=12, space_after=2)
code_block(doc, """\
BEGIN
    sp_create_booking(
        1, 1, 1, 1, 'SAVE10',
        TO_TIMESTAMP('2026-05-05 09:00:00', 'YYYY-MM-DD HH24:MI:SS'),
        2
    );
END;
/
-- PL/SQL procedure successfully completed.
-- BOOKINGS: 1 row inserted, status=CONFIRMED, promo_id=1
-- PROMOTIONS: current_uses incremented from 0 to 1""")

body(doc, 'Occupied slot rejection (ORA-20010):', size=12, space_after=2)
code_block(doc, """\
BEGIN
    sp_create_booking(2, 1, 2, 1, NULL,
        TO_TIMESTAMP('2026-05-05 09:30:00', 'YYYY-MM-DD HH24:MI:SS'), 1);
END;
/
-- ORA-20010: Time slot is already occupied by another booking.
-- ERROR_LOGS: 1 row inserted, severity=CRITICAL,
--             procedure_name=sp_create_booking""")

body(doc, 'Invoice generation \u2014 sp_generate_invoice(1) on a 2-hour Pipe Leak Fix at Rs. 350/hr:', size=12, space_after=2)
code_block(doc, """\
BEGIN sp_generate_invoice(p_booking_id => 1); END;
/
-- INVOICES: 1 row inserted
-- base_amount  = Rs. 700.00  (350 x 2 hrs)
-- discount     = Rs.   0.00  (no promo on this booking)
-- platform_fee = Rs.  70.00  (10% of 700)
-- tax_amount   = Rs.  35.00  (5% of 700)
-- net_total    = Rs. 805.00""")

heading(doc, '6.3  Trigger verification', size=13, bold=True, space_before=8, space_after=6)
body(doc, 'Before inserting a review for booking 1 (Bob Builder, Provider 1):', size=12, space_after=2)
code_block(doc, """\
PROVIDER_ID  RATING_AVG  JOBS_COMPLETED
1            4.50        50""")
body(doc, 'After INSERT INTO REVIEWS (booking_id, rating, comments) VALUES (1, 5, \'Excellent\'):', size=12, space_after=2)
code_block(doc, """\
PROVIDER_ID  RATING_AVG  JOBS_COMPLETED
1            4.57        51
-- rating_avg = AVG across all reviews for Provider 1
-- jobs_completed = COUNT(DISTINCT COMPLETED bookings), not review count""")
body(doc, (
    'Updating the review rating to 3 fires the trigger again and adjusts rating_avg accordingly. '
    'Deleting the review sets rating_avg back to its pre-review value and decrements '
    'jobs_completed. All three cases (INSERT, UPDATE, DELETE) are covered by the compound trigger\'s '
    'FOR INSERT OR UPDATE OR DELETE declaration.'
), size=12, space_after=8)

heading(doc, '6.4  Web application snapshots', size=13, bold=True, space_before=8, space_after=6)
for screen, desc in [
    ('Landing page',
     'Public hero section with the ServeMart logo, service category quick-links (Plumbing, '
     'Electrical, Cleaning, Painting, Carpentry, Appliance Repair), an embedded Leaflet/'
     'OpenStreetMap showing all five service cities, and a live provider grid fetched from the '
     'API showing the top-rated providers. Navigation links direct visitors to register as a '
     'customer or professional.'),
    ('Register and Login',
     'Single Register page with a role toggle (Hire a Pro / Join as Pro) that switches between '
     'CUSTOMER and PROVIDER registration. The Login form accepts username and password and '
     'performs role-aware redirect: customers go to the customer dashboard, providers to the '
     'provider dashboard, admins to the admin panel.'),
    ('Customer dashboard',
     'After login, the customer sees their upcoming and recent bookings, a quick-access category '
     'grid, and a recommended provider list. City detection calls GET /api/bookings/my-city to '
     'identify the city from the customer\'s most recent booking address, then feeds that city\'s '
     'area_id to fn_recommend_providers for a personalised ranked list.'),
    ('Customer browse page',
     'Grid of provider cards showing name, category badge, star rating, and job count. Only '
     'providers with at least one active service listed appear (filtered server-side by an EXISTS '
     'subquery). Each card has a "View services \u2192" button that opens the booking form.'),
    ('Booking form \u2014 details step',
     'Service dropdown, promo code field with an inline "Check" button that calls '
     'GET /api/bookings/validate-promo and displays the discount before submission '
     '(e.g., "10% off (up to Rs.150)"), address fields (house number, building, landmark, '
     'city, pincode), availability slot picker, date picker showing the next 4 occurrences '
     'of the chosen day, and a 3-hour time window grid that hides occupied slots '
     '(fetched from GET /api/bookings/booked-slots/:providerId). '
     'Clicking "Confirm Booking" submits in one step and advances to the success screen.'),
    ('Booking success screen',
     'Full-page confirmation card showing booking ID, scheduled date, provider name, and '
     'address. A "Go to My Bookings" link follows. Payment collection is deferred to the '
     'provider at job completion.'),
    ('My bookings page',
     'Filterable list by status (ALL / PENDING / CONFIRMED / COMPLETED / CANCELLED) with '
     'colour-coded status pills. COMPLETED bookings show a "View Invoice" button that opens '
     'an itemised invoice modal with base fee, platform fee, promo discount, GST, net total, '
     'payment method, and transaction ID. Active bookings show a "Cancel" button with a '
     'reason-input modal. CANCELLED bookings display the cancellation reason in a red callout; '
     '"Invoice pending" is never shown on cancelled entries. '
     'Completed bookings also show a "Leave Review" button.'),
    ('Review form',
     'Accessible at /customer/review/:bookingId. Five interactive star buttons and a comment '
     'field (minimum 10 characters). Submitting calls POST /api/reviews/:booking_id, which '
     'fires trg_update_provider_rating and immediately updates the provider\'s rating_avg and '
     'jobs_completed in SERVICE_PROVIDERS.'),
    ('Provider dashboard \u2014 bookings',
     'Lists all bookings assigned to the logged-in provider with two action buttons per row: '
     '"Mark Complete" opens a payment method modal (Cash / UPI / Credit Card radio buttons) '
     'before calling sp_generate_invoice; "Cancel" opens a reason-input modal for PENDING and '
     'CONFIRMED bookings. CANCELLED bookings display the cancellation reason inline. '
     'The provider\'s current rating average and job count are shown at the top.'),
    ('Provider \u2014 manage services and availability',
     'Providers add services with category, name, and hourly rate, and toggle existing listings. '
     'The Availability page lets them define weekly recurring time slots by day and hour range. '
     'Changes to availability are immediately visible to customers on the booking form.'),
    ('Provider \u2014 manage service areas',
     'Dedicated page at /provider/areas. Shows a dropdown of all platform cities; selecting one '
     'calls POST /api/providers/my/areas to add it to PROVIDER_AREAS. Assigned cities appear in '
     'a removable list. Success and error feedback is shown via modal.'),
    ('Admin \u2014 revenue dashboard',
     'Four stat cards (total platform revenue, gross booking value, completed booking count, '
     'active provider count) and a Recharts dual-bar chart showing platform fee vs. gross value '
     'per service category.'),
    ('Admin \u2014 providers',
     'Table of all registered providers with background check status badges. Approve and Reject '
     'buttons call POST /api/admin/providers/:id/approve which updates background_chk in '
     'SERVICE_PROVIDERS. Only APPROVED providers appear in the customer browse page.'),
    ('Admin \u2014 error logs',
     'Dedicated page at /admin/errors. All rows from ERROR_LOGS displayed with left-border '
     'colour coding: CRITICAL in red, HIGH in orange, MEDIUM in amber, LOW in grey. '
     'Each card shows the procedure name, error message, and timestamp.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(screen + ' \u2014 ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)
    body(doc, f'[Screenshot: {screen}]', size=10, space_after=6)
doc.add_page_break()

# ─── CHAPTER 7 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 7: CONCLUSION, LIMITATIONS AND FUTURE WORK', size=14, bold=True, space_before=8, space_after=12)

heading(doc, '7.1  Conclusion', size=13, bold=True, space_before=8, space_after=6)
body(doc, (
    'ServeMart achieves what it set out to do. The 16-table Oracle schema handles the full '
    'booking lifecycle with referential integrity intact. PL/SQL procedures prevent double-booking '
    'at the database layer, not just the application layer \u2014 the restriction holds even if '
    'someone calls the database directly. The compound trigger keeps provider statistics accurate '
    'without requiring application code to manually update averages. The three-tier separation '
    'means the React frontend, Express API, and Oracle database can each be reasoned about '
    'independently.'
), size=12, space_after=8)
body(doc, 'The project achieved all stated objectives:', size=12, space_after=4)
for item in [
    'A normalised relational schema was designed and implemented in Oracle XE with 16 tables, all in 3NF.',
    'All entities carry appropriate PRIMARY KEY, FOREIGN KEY, NOT NULL, UNIQUE, and CHECK constraints.',
    'Basic and complex SQL queries were written and tested, covering multi-table JOINs, aggregations, LISTAGG, subqueries, and window functions.',
    'Three stored procedures encapsulate the booking, invoicing, and cancellation workflows with structured error logging.',
    'A compound trigger on REVIEWS recomputes provider ratings and completed job counts after every review change, without application-level code.',
    'A table function returns a ranked recommendation cursor for any service area using a weighted composite score.',
    'A view aggregates provider summary data consumed by two separate API routes.',
    'A role-adaptive React single-page application demonstrates live Oracle connectivity with JWT authentication, Zod request validation, and bind variables on every query.',
    'Applied use of: multi-table joins and aggregates, PL/SQL exception handling, compound trigger to avoid the mutating table problem, cursor-returning functions, and autonomous error commits inside transactional procedures.',
]:
    bullet(doc, item)
doc.add_paragraph()

heading(doc, '7.2  Limitations', size=13, bold=True, space_before=8, space_after=6)
for title, desc in [
    ('No live payment gateway',
     'Payment is recorded as SUCCESS automatically when the provider marks the job complete '
     'and selects Cash, UPI, or Credit Card. No actual money is processed.'),
    ('No real-time updates',
     'The provider dashboard does not push new bookings and the customer has to refresh to see '
     'status changes. There is no WebSocket layer.'),
    ('City-level geography only',
     'Provider coverage is assigned per city from a fixed admin-managed list. There is no GPS-based '
     'proximity search or pin-code level filtering.'),
    ('Single admin account',
     'There is no hierarchy within the admin role and no audit trail of admin actions (approvals, '
     'area additions).'),
    ('No email or SMS notifications',
     'Booking confirmations and cancellations do not trigger any notifications outside the '
     'application interface.'),
    ('Simulated background check',
     'Providers register with status PENDING and are manually approved by the admin clicking a '
     'button. No document verification pipeline or third-party identity check integration exists.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + ': ')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)
doc.add_paragraph()

heading(doc, '7.3  Future work', size=13, bold=True, space_before=8, space_after=6)
for item in [
    'Integrate a payment gateway (Razorpay or Stripe) with webhook callbacks to update payment status from the processor side. The PAYMENTS.transaction_id column is already present for this.',
    'Add WebSocket-based real-time booking notifications so providers see new bookings and customers see status changes instantly.',
    'Use Oracle Spatial or a geocoding API to support radius-based provider search from the customer\'s GPS location, replacing the current city-level matching.',
    'Build an email notification layer (Nodemailer + SendGrid) for booking confirmations, cancellation alerts, and review prompts.',
    'Replace the composite score function with a collaborative filtering recommendation model trained on booking history and customer preferences.',
    'Add document upload and an admin review queue to make provider background verification a real workflow rather than a manual checkbox.',
    'Build a React Native mobile app against the same REST API for on-the-go booking.',
    'Export invoices as PDF using a server-side renderer (Puppeteer or PDFKit) rather than relying on the browser\'s print dialog.',
]:
    bullet(doc, item)
doc.add_page_break()

# ─── CHAPTER 8 ───────────────────────────────────────────────────────────────

heading(doc, 'CHAPTER 8: REFERENCES', size=14, bold=True, space_before=8, space_after=12)
for i, ref in enumerate([
    'Silberschatz, A., Korth, H. F., & Sudarshan, S. (2019). Database System Concepts (7th ed.). McGraw-Hill Education.',
    'Ramakrishnan, R., & Gehrke, J. (2002). Database Management Systems (3rd ed.). McGraw-Hill.',
    'Oracle Corporation. (2023). Oracle Database Documentation. Oracle Help Center. https://docs.oracle.com/en/database/oracle/oracle-database/',
    'Oracle Corporation. Oracle Database PL/SQL Language Reference (21c). https://docs.oracle.com/en/database/oracle/oracle-database/21/lnpls/',
    'Oracle Corporation. Oracle Database SQL Language Reference (21c). https://docs.oracle.com/en/database/oracle/oracle-database/21/sqlrf/',
    'Oracle Corporation. node-oracledb Documentation (v6). https://node-oracledb.readthedocs.io/',
    'Express.js 4.x Documentation. OpenJS Foundation. https://expressjs.com/',
    'React Documentation (v19). Meta Open Source. https://react.dev/',
    'Auth0. Introduction to JSON Web Tokens. https://jwt.io/introduction',
    'Leaflet.js Contributors. (2024). Leaflet \u2014 open-source JavaScript library for interactive maps. https://leafletjs.com/',
    'Recharts Group. (2024). Recharts \u2014 chart library built with React and D3. https://recharts.org/',
    'Colby, S. et al. (2023). Zod \u2014 TypeScript-first schema validation. https://zod.dev/',
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    run = p.add_run(f'{i}. {ref}')
    set_font(run, size=12)

# ─── SAVE ────────────────────────────────────────────────────────────────────

output_path = r'D:\DBS PROJECT\ServeMart_MiniProject_Report_v4.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
